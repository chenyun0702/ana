import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import json
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx_report():
    dataset_path = r'C:\Users\927632_st.tc\.gemini\antigravity\scratch\gy61-vibration-analysis\parsed_dataset.json'
    with open(dataset_path, 'r', encoding='utf-8') as f:
        ds = json.load(f)

    stats_list = ds['summary_stats']

    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = '微軟正黑體'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Title & Subtitle
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("基於 GY-61 加速度感測器與調諧質量阻尼器 (TMD)\n之高樓減震效能實驗研究報告")
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0xEC, 0x48, 0x99)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("臺中市 114 學年度國民教育階段資賦優異學生獨立研究專題成果\n研究作者：張玲瑀、陳姿妘")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_paragraph()

    # Executive Summary Box
    table_callout = doc.add_table(rows=1, cols=1)
    table_callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_c = table_callout.cell(0, 0)
    set_cell_background(cell_c, "FCE7F3")
    set_cell_margins(cell_c, top=150, bottom=150, left=200, right=200)

    p_call = cell_c.paragraphs[0]
    p_call.paragraph_format.space_before = Pt(4)
    p_call.paragraph_format.space_after = Pt(4)
    run_call_t = p_call.add_run("✨ 摘要與核心成果總結 (Executive Summary)\n")
    run_call_t.bold = True
    run_call_t.font.size = Pt(12)
    run_call_t.font.color.rgb = RGBColor(0xBE, 0x18, 0x5D)

    summary_text = (
        "本研究旨在探討調諧質量阻尼器 (Tuned Mass Damper, TMD / 減震球) 於高樓建築抗震中之減震效能。 "
        "實驗採用 3D 列印高樓模型、MediaTek LinkIt 7697 / ESP32 Node32s 物聯網控制板、EZ Start Kit 轉速控制馬達震動台， "
        "並以 GY-61 (ADXL335) 三軸加速度感測器進行動態訊號擷取，透過 Bluetooth Terminal 進行無線數據傳輸。\n\n"
        "本次實驗共完成 3 種擺重 (10g, 20g, 30g)、3 種擺繩長度 (5cm, 10cm, 15cm) 與 3 種震動速度 (慢速、中速、快速) 之 27 組獨立條件測試， "
        "總計分析 271 筆時序加速度數據。實驗結果顯示：\n"
        "1. 擺長效應：擺繩長度達 15cm 時，單擺自然頻率 (1.29 Hz) 與高樓共振晃動週期達成精準調諧匹配，減震效益最為顯著。\n"
        "2. 擺重效應：阻尼球重量並非越重越好。適中質量 (10g~20g) 能有效吸收能量；過重 (30g) 會產生拉扯樓體之二次諧振現象。\n"
        "3. 最佳抗震組合：在馬達快速震動強迫共震條件下，「15cm 擺長 + 10g 擺重」達成最佳抗震配置，將頂樓晃動 Peak-to-Peak 振幅由基準的 9.20g 大幅降低至 1.2518g，達成高達 86.39% 的減震抑制率！"
    )
    run_call_b = p_call.add_run(summary_text)
    run_call_b.font.size = Pt(10.5)

    doc.add_paragraph()

    # Chapter 1
    h1 = doc.add_heading('第一章 緒論', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_heading('1.1 研究背景與動機', level=2)
    doc.add_paragraph(
        "台灣地位於環太平洋地震帶，地震頻繁且規模巨大（如 1935 年新竹台中地震、1999 年 921 大地震、2016 年美濃地震及 2024 年花蓮地震）。 "
        "隨著現代都市高樓大廈林立，結構在地震與強風作用下易產生劇烈共振晃動，造成內部人員不適與結構損傷。 "
        "常見的高樓抗震技術中，「調諧質量阻尼器 (TMD)」透過於建築頂樓懸掛大質量重錘，利用運動相位差反向抵銷結構擺動能量。 "
        "本研究希望透過微型感測器與控制板，親自設計實驗驗證 TMD 在不同結構與外力條件下的減震物理規律。"
    )

    doc.add_heading('1.2 研究目的', level=2)
    doc.add_paragraph(
        "1. 結合 LinkIt 7697 / Node32s 與 GY-61 (ADXL335) 加速度傳感器，建構高樓模型即時震動監測系統。\n"
        "2. 定量分析阻尼球質量 (10g, 20g, 30g) 對頂樓搖晃幅度的影響。\n"
        "3. 定量分析阻尼擺繩長度 (5cm, 10cm, 15cm) 對共震調諧匹配度的影響。\n"
        "4. 評估不同震動頻率（慢速、中速、快速）下，無阻尼器與有阻尼器配置之減震抑制率 (Vibration Reduction %)。"
    )

    # Chapter 2
    h2 = doc.add_heading('第二章 實驗原理與儀器設備', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_heading('2.1 調諧質量阻尼器 (TMD) 物理原理', level=2)
    doc.add_paragraph(
        "阻尼擺錘懸掛於頂樓，可簡化為單擺系統。其理論自然振動頻率 f_TMD (Hz) 公式如下：\n"
        "f_TMD = (1 / 2π) * √(g / L)\n"
        "其中 g = 9.81 m/s²，L 為擺繩長度 (m)。\n"
        "• 當 L = 5 cm (0.05m) 時，理論頻率 f_TMD ≈ 2.23 Hz\n"
        "• 當 L = 10 cm (0.10m) 時，理論頻率 f_TMD ≈ 1.58 Hz\n"
        "• 當 L = 15 cm (0.15m) 時，理論頻率 f_TMD ≈ 1.29 Hz"
    )

    doc.add_heading('2.2 GY-61 (ADXL335) 加速度傳感器訊號校正', level=2)
    doc.add_paragraph(
        "GY-61 採用 Analog Devices ADXL335 晶片，為三軸類比電壓輸出加速度計。 "
        "感測器量測範圍為 ±3g，靈敏度 S = 300 mV/g，靜止 Zero-G 參考電壓 V_zero ≈ 1.65V。 "
        "LinkIt 7697 / ESP32 提供 12-bit 類比數位轉換器 (ADC, 0~4095 count)。轉換公式為：\n"
        "a_y = [(ADC_y - 2048) / 4095] * (3.3V / 0.3V/g)\n"
        "峰對峰振幅 (Peak-to-Peak Amplitude): A_p-p = a_max - a_min (g)\n"
        "均方根加速度 (RMS): a_RMS = √( (1/N) * Σ a_i² ) (g)"
    )

    # Chapter 3
    h3 = doc.add_heading('第三章 實驗數據與全條件統計結果', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_paragraph("以下為全套 27 組獨立實驗條件之 GY-61 ADXL335 加速度實測統計數據：")

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    hdr_titles = ["組別代碼", "擺重 (g)", "擺長 (cm)", "馬達轉速", "Y軸振幅 Ap-p (g)", "RMS 加速度 (g)", "減震抑制率 (%)"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0284C7")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)

    for stats in stats_list:
        row_cells = table.add_row().cells
        vals = [
            stats['cond_id'],
            str(stats['weight_g']),
            f"{stats['length_cm']:.0f}",
            stats['speed'],
            f"{stats['peak_to_peak_y']:.4f}",
            f"{stats['rms_y']:.4f}",
            f"{stats['vibration_reduction_pct']:.2f}%"
        ]
        
        is_best = (stats['cond_id'] == 'w10_l15cm_快')
        bg_color = "FCE7F3" if is_best else "FFFFFF"

        for i, val in enumerate(vals):
            row_cells[i].text = val
            set_cell_background(row_cells[i], bg_color)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.size = Pt(9)
                if is_best:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xBE, 0x18, 0x5D)

    doc.add_paragraph()

    # Chapter 4
    h4 = doc.add_heading('第四章 數據分析與結果討論', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_heading('4.1 擺長效應 (Length Effect) 分析', level=2)
    doc.add_paragraph(
        "對比 5cm、10cm 與 15cm 擺繩長度在快速晃動下的表現：\n"
        "• 5cm 擺長：平均 Peak-to-Peak 振幅為 5.08g，減震率約 44.7%。擺繩過短導致單擺自然頻率 (2.23Hz) 偏高，無法配合高樓主要搖晃頻率。\n"
        "• 10cm 擺長：平均 Peak-to-Peak 振幅為 4.54g，減震率約 50.5%。\n"
        "• 15cm 擺長：平均 Peak-to-Peak 振幅大幅降至 2.26g，平均減震率高達 75.4%！\n"
        "實驗數據證實：擺繩長度較長 (15cm) 時，阻尼器週期與高樓樓體搖晃週期達到最佳匹配與相位共振抑制。"
    )

    doc.add_heading('4.2 擺重效應 (Mass Effect) 分析', level=2)
    doc.add_paragraph(
        "對比 10g、20g 與 30g 阻尼球質量對頂樓晃動的影響：\n"
        "• 10g 擺重：在 15cm 擺長下發揮極佳的反相位慣性抑制，振幅最低 (1.2518g)。\n"
        "• 20g 擺重：減震效果次之 (2.4015g)。\n"
        "• 30g 擺重：頂樓振幅回升至 3.1348g。\n"
        "實驗結果證實：質量並非越重越好。當擺重過重 (30g) 時，擺錘甩動產生的慣性力反向拉扯樓體，引發二次諧振晃動；適中質量 (10g~20g) 方為最佳解。"
    )

    doc.add_heading('4.3 最佳抗震配置排行榜 (Top 5 Optimal Configurations)', level=2)
    
    sorted_stats = sorted(stats_list, key=lambda x: x['peak_to_peak_y'])[:5]

    table_top = doc.add_table(rows=1, cols=5)
    table_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_top)

    t_hdr = table_top.rows[0].cells
    t_titles = ["名次", "實驗條件組別", "Y軸 Peak-to-Peak (g)", "RMS 加速度 (g)", "減震抑制率 (%)"]
    for i, title in enumerate(t_titles):
        t_hdr[i].text = title
        set_cell_background(t_hdr[i], "EC4899")
        p = t_hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)

    for rank, st in enumerate(sorted_stats, 1):
        row_c = table_top.add_row().cells
        row_c[0].text = f"第 {rank} 名"
        row_c[1].text = st['cond_id']
        row_c[2].text = f"{st['peak_to_peak_y']:.4f} g"
        row_c[3].text = f"{st['rms_y']:.4f} g"
        row_c[4].text = f"{st['vibration_reduction_pct']:.2f}%"

        bg = "FFF1F2" if rank == 1 else "FFFFFF"
        for idx in range(5):
            set_cell_background(row_c[idx], bg)
            p = row_c[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.size = Pt(9)
                if rank == 1:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xBE, 0x18, 0x5D)

    doc.add_paragraph()

    # Chapter 5
    h5 = doc.add_heading('第五章 結論與建議', level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_heading('5.1 主要研究結論', level=2)
    doc.add_paragraph(
        "1. 本獨立研究成功結合 LinkIt 7697 / Node32s 與 GY-61 (ADXL335) 加速度傳感器，建立了低成本、高精度的地震晃動量測與分析平台。\n"
        "2. 擺長為決定 TMD 共震調諧的核心變因。擺繩長度長 (15cm) 之自然頻率 (1.29Hz) 能精準匹配高樓共振頻率，減震表現最突出。\n"
        "3. 阻尼球質量需精細調諧。過重阻尼球 (30g) 會產生二次反拉力；在本實驗模型中，「15cm 擺長 + 10g 擺重」達成 86.39% 最高減震率。"
    )

    doc.add_heading('5.2 未來改進與延伸應用', level=2)
    doc.add_paragraph(
        "• 增加主動式電磁驅動控制 (Active Mass Damper, AMD)，由 GY-61 即時反饋訊號主動控制馬達輸出反向抵銷力。\n"
        "• 增加雙軸與三軸複合阻尼球，測試非正交方向地震波之三維減震效能。"
    )

    doc.add_heading('5.3 參考文獻', level=2)
    doc.add_paragraph(
        "1. 台灣物聯科技 GY-61 ADXL335 三軸加速度感測器規格書 (https://www.taiwaniot.com.tw/)\n"
        "2. 宇倉電子有限公司 旋轉電位器與 PWM 馬達驅動板教學手冊\n"
        "3. Analog Devices, ADXL335 Small, Low Power, 3-Axis ±3g Accelerometer Datasheet.\n"
        "4. MediaTek LinkIt 7697 / ESP32 物聯網開發板官方說明文件。"
    )

    target_paths = [
        r'C:\Users\927632_st.tc\Documents\減震球\阻尼器加速度分析網頁\GY61_阻尼器減震實驗研究報告.docx',
        r'C:\Users\927632_st.tc\Documents\減震球\阻尼器加速度分析網頁\GY61_阻尼器減震實驗研究報告_最新版.docx',
        r'C:\Users\927632_st.tc\.gemini\antigravity\scratch\gy61-vibration-analysis\GY61_阻尼器減震實驗研究報告.docx'
    ]

    for p in target_paths:
        try:
            doc.save(p)
            print(f"Successfully saved Word report to: {p}")
        except Exception as e:
            print(f"Could not save to {p}: {e}")

if __name__ == "__main__":
    build_docx_report()
